from dotenv import load_dotenv
import smtplib, time, os, openpyxl,requests,re
from email.header import Header
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from bs4 import BeautifulSoup
from datetime import datetime
import feedparser
import pandas as pd
from collections import Counter
from docx import Document
from docx2pdf import convert
from openpyxl import Workbook
import json
from pymongo import MongoClient
import pythoncom

class Criminal:
    def __init__(self, ip_address):
        
        self.url = f"https://api.criminalip.io/v1/asset/ip/summary?ip={ip_address}"

        self.payload={}
        self.headers = {
            "x-api-key": "<YOUR Criminal API KEY>"
        }
    def get_criminal_info(self):
        response = requests.request("GET", self.url, headers=self.headers, data=self.payload)
        if response.status_code == 200:
            AE = AutoExcel()
            print(response.json())
            AE.criminal_info_excel1(response.json())
            return response.json()
        else:
            print("접속 오류 발생")


    def printing(result):
        print("IP 정보:")
        print(f"IP 주소: {result['ip']}")
        print(f"점수 (수신): {result['score']['inbound']}")
        print(f"점수 (발신): {result['score']['outbound']}")
        print(f"국가: {result['country']}")
        print(f"국가 코드: {result['country_code']}")
        print(f"지역: {result['region']}")
        print(f"도시: {result['city']}")
        print(f"ISP: {result['isp']}")
        print(f"기관 이름: {result['org_name']}")
        print(f"AS 번호: {result['as_no']}")
        print(f"우편 코드: {result['postal_code']}")
        print(f"위도: {result['latitude']}")
        print(f"경도: {result['longitude']}")
    
class File:
    def __init__(self,file,filename):
        self.file = file
        self.filename = filename
        self.df = None
        self.temp_file_path = None
        self.today = datetime.now().strftime('%Y-%m-%d')
        self.timestamp = datetime.now().strftime('%Y-%m-%d_%H-%M')
        self.copy_file = None
        self.content = ""
        
    def extension(self):
        if self.filename.endswith(".txt"):
            self.temp_file_path = os.path.join('uploads',f'{self.timestamp}_{self.filename}_copy.csv')
            self.read_text()
        elif self.filename.endswith(".xlsx"):
            self.temp_file_path = os.path.join('uploads',f'{self.timestamp}_{self.filename}_copy.xlsx')
            self.read_excel()
        elif self.filename.endswith(".docx"):
            self.temp_file_path = os.path.join('uploads',f'{self.timestamp}_{self.filename}_copy.docx')
            self.read_docx()
        elif self.filename.endswith(".log"):
            self.temp_file_path = os.path.join('uploads',f'{self.timestamp}_{self.filename}_copy.log')
            self.read_log()
    
    def read_text(self):
        self.file.save(self.temp_file_path)
        self.df = pd.read_csv(self.temp_file_path, encoding='utf-8')
        return self.df

    def read_excel(self):
        self.file.save(self.temp_file_path)
        self.df = pd.read_excel(self.temp_file_path)
        if self.find_email():
            mail = Mail("hanmin9981@naver.com")
            mail.mail_text_sender(self.temp_file_path, "\n email 포함된 파일이 업로드 되었습니다.")
        return self.df

    def find_email(self):
        files = []
        contain_symbol = self.df.apply(lambda col: col.apply(self.contains_at_symbol))
        print("@ 포함")
        return any(contain_symbol)

    
    def contains_at_symbol(self, cell):
        cell_str = str(cell) if pd.notnull(cell) else ''
        return bool(re.search(r"[\w\.-]+@[\w\.-]+", cell_str))

    def read_docx(self):
        self.file.save(self.temp_file_path)
        self.df = Document(self.temp_file_path)
        for paragraph in self.df.paragraphs:
            if 'NAME' in paragraph.text:
                paragraph.text = paragraph.text.replace('NAME', '조정원')
            elif 'FILE' in paragraph.text:
                paragraph.text = paragraph.text.replace('FILE', f'{self.filename}')
            elif 'DATE' in paragraph.text:
                paragraph.text = paragraph.text.replace('DATE', '200')
            elif 'INFO' in paragraph.text:
                paragraph.text = paragraph.text.replace('COUNT', '200')

        doc_file = f'{self.temp_file_path}01.docx'
        pdf_file = f'{self.temp_file_path}01.pdf'
        self.df.save(doc_file)
        convert(doc_file, pdf_file)
        return self.df

    def read_log(self):
        print(self.temp_file_path)
        self.file.save(self.temp_file_path)
        result_list = []
        ip_list = []
        timestamp_list = []
        with open(self.temp_file_path, 'r') as f:
            for line in f:
                match = re.findall(r"\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}", line)
                if match:
                    ip_list.append(match[0])
                    timestamp_match = re.findall(r"\d{4}-\d{2}-\d{2} \d{2}:\d{2}", line)
                if timestamp_match:
                    timestamp_list.append(timestamp_match[0])
        print("log : ", ip_list[0]) # for문으로 10개 IP CRIMINAL로 데이터화 하기.
        info = Criminal(ip_list[0]).get_criminal_info()
        ip_counter = Counter(ip_list)
        top_10_ips = ip_counter.most_common(10)
        self.copy_file = f'{self.timestamp}_{self.filename}_copy.txt'


        with open(os.path.join('uploads',self.copy_file), 'w', encoding='utf-8') as f:
            for ip_address, count in top_10_ips:
                log_timestamps = []
                for ip, timestamp in zip(ip_list, timestamp_list):
                    if ip == ip_address:
                        log_timestamps.append(timestamp)
                if log_timestamps:
                    latest_timestamp = max(log_timestamps)
                else:
                    latest_timestamp = ''
                time_counter = Counter()
                for timestamp in log_timestamps:
                    time_counter[timestamp] += 1
                time_counts = time_counter.most_common()

                result_list.append(f"IP 주소: {ip_address}\n")
                for timestamp, time_count in sorted(time_counts):
                    result_list.append(f"시간: {timestamp}, 접속 횟수: {time_count}\n")
                result_list.append(f"총 접속 횟수: {count}, 최근 접속 시간: {latest_timestamp}\n")
                result_list.append("\n")
                f.writelines(result_list)

                #line = f"{ip_address} 접속 횟수 {count}, 가장 최근 접속 시간: {str(latest_timestamp)}\n"
                
                #f.write(line)
        self.create_report(info)
    
    def create_report(self,info):
        doc = Document(os.path.join('uploads','template.docx'))
        file_path = os.path.join('uploads', self.copy_file)
        DB = DBConnect()
        DB.inData(self.filename, datetime.now())
        print("DB 삽입 성공")
        # 텍스트 파일을 읽어 리스트로 변환
        with open(file_path, 'r', encoding='utf-8') as file:
            all_lines = file.readlines()

        # 각 행의 끝에 있는 개행 문자('\n') 제거
        all_lines = [line.strip() for line in all_lines]
        print("list : ",all_lines)
        for paragraph in doc.paragraphs:
            if 'NAME' in paragraph.text:
                paragraph.text = paragraph.text.replace('NAME', '4조')
            elif 'DATE' in paragraph.text:
                paragraph.text = paragraph.text.replace('DATE', f'{self.today}')
            elif 'CONTENT' in paragraph.text:
               for con in all_lines:
                   self.content+="\t"+con+"\n"
               paragraph.text = paragraph.text.replace('CONTENT', f'{self.content}')
            elif 'FILE' in paragraph.text:
               paragraph.text = paragraph.text.replace('FILE', f'{self.filename}')
            elif 'INFO' in paragraph.text:
                paragraph.text = paragraph.text.replace('INFO', f'{info}')

        doc_file = f'{self.timestamp}_{self.filename}_report.docx'
        pdf_file = f'{self.timestamp}_{self.filename}_report.pdf'
        doc.save(doc_file)
        

        pythoncom.CoInitialize()
        convert(doc_file, pdf_file)

class AutoExcel:
    def __init__(self):
        self.today = datetime.now().strftime('%Y-%m-%d')
        self.workbook = None
        self.worksheet = None
        self.header = {
            'User-Agent': 'Mozilla/5.0',
            'Content-Type': 'text/html; charset=utf-8'
        }
        self.timestamp = datetime.now().strftime('%Y-%m-%d_%H-%M')

    def criminal_info_excel(self, result):
        self.workbook = Workbook()
        self.worksheet = self.workbook.active
        #json_data = json.load(result)
        for row_index, (key, value) in enumerate(result.items(), start=1):
            self.worksheet.cell(row=row_index, column=1, value=key)
            if key == 'score' and isinstance(value, dict):
                value = json.dumps(value)

            self.worksheet.cell(row=row_index, column=2, value=value)

        file_path = os.path.join('uploads', f'{self.timestamp}_ip_info.xlsx')
        self.workbook.save(file_path)

    def criminal_info_excel1(self, result):
        try:
            self.workbook = Workbook()
            self.worksheet = self.workbook.active

            for row_index, (key, value) in enumerate(result.items(), start=1):
                self.worksheet.cell(row=row_index, column=1, value=key)
                if key == 'score' and isinstance(value, dict):
                    value = json.dumps(value)
                if isinstance(value, dict) and not value:
                    value = None  # 빈 딕셔너리는 None으로 대체
                self.worksheet.cell(row=row_index, column=2, value=value)

            file_path = os.path.join('uploads', f'{self.today}_ip_info.xlsx')
            self.workbook.save(file_path)
            print(f"Excel file saved successfully.")
        except Exception as e:
            print(f"Error saving Excel file: {e}")




class DirControl:
    def __init__(self, dir_path):
        self.file_path = os.path.join(dir_path)
        self.all_files = os.listdir(self.file_path)
        self.text_files = []
        self.py_files = []
        self.zip_files = []
        self.xlsx_files = []
        self.extension_files = []

    def search_text_files(self):
        for file in self.all_files:
            if file.endswith('txt'):
                self.text_files.append(file)
        
        return self.text_files

    def search_py_files(self):
        for file in self.all_files:
            if file.endswith('py'):
                self.py_files.append(file)
        
        return self.py_files
    
    def search_zip_files(self):
        for file in self.all_files:
            if file.endswith('zip'):
                self.zip_files.append(file)
        
        return self.zip_files

    def search_xlsx_files(self):
        for file in self.all_files:
            if file.endswith('xlsx'):
                self.xlsx_files.append(file)
        
        return self.xlsx_files  
    
    def search_files(self,extension):
        for file in self.all_files:
            if file.endswith(extension):
                self.extension_files.append(file)
        
        return self.extension_files  

class FileControl:
    def __init__(self):
        pass



class Mail:
    def __init__(self,received_email):
        load_dotenv()
        self.SECRET_ID = os.getenv("SECRET_ID")
        self.SECRET_PASS = os.getenv("SECRET_PASS")
        self.smtp = None
        self.recv_email = received_email
        
    def settings(self):
        self.smtp = smtplib.SMTP('smtp.naver.com', 587)
        self.smtp.ehlo()
        self.smtp.starttls()
        self.smtp.login(self.SECRET_ID, self.SECRET_PASS)
        

    def mail_text_sender(self, file_path, msg_list):
        self.settings()
        myemail = f"{self.SECRET_ID}@naver.com"
        youremail = self.recv_email
        subject = f"{file_path} 파일 탐지"
        message = f"탐지된 파일 : {file_path}"
        for msg in msg_list:
            message += msg
        msg = MIMEText(message.encode('utf-8'), _subtype='plain', _charset='utf-8')
        msg['Subject'] = Header(subject.encode('utf-8'), 'utf-8')
        msg['From'] = myemail
        msg['To'] = youremail
        self.smtp.sendmail(myemail,youremail,msg.as_string())
        self.smtp.quit()
    
    def mail_attach_sender(self,file_path,send_file):
        self.settings()
        myemail = f"{self.SECRET_ID}@naver.com"
        youremail = self.recv_email

        msg = MIMEMultipart()

        msg['Subject'] ="첨부파일 입니다."
        msg['From'] = myemail
        msg['To'] = youremail

        text = """
        첨부파일 메일 입니다.
        감사합니다.
        """
        contentPart = MIMEText(text) 
        msg.attach(contentPart) 

        etc_file_path = os.path.join(file_path,send_file)
        with open(etc_file_path, 'rb') as f : 
            etc_part = MIMEApplication( f.read() )
            etc_part.add_header('Content-Disposition','attachment', filename=etc_file_path)
            msg.attach(etc_part)

        self.smtp.sendmail( myemail,youremail,msg.as_string() )
        self.smtp.quit()


class Monitering:
    def __init__(self,DIR_PATH,received_email):
        self.DIR_WATCH = DIR_PATH
        self.previous_files = set(os.listdir(DIR_PATH))
        self.inspect_list = []
        self.received_email = received_email
        self.mail = Mail(self.received_email)

    def inspect_annotation(self, word):
        while True:
            time.sleep(1)
            self.inspect_list = []
            current_files= set(os.listdir(self.DIR_WATCH))
            new_files = current_files - self.previous_files
            for filename in new_files:
                if filename.endswith(".txt") or filename.endswith(".csv"):
                    file_path = os.path.join(self.DIR_WATCH, filename)
                    with open(file_path,'r', encoding='utf-8') as f:
                        lines = f.readlines()
                        for line in lines:
                            if line.startswith(f"{word}") or line.startswith("//") or line.startswith("#"):
                                self.inspect_list.append(f"{line}")
                
                    if self.inspect_list:
                        print(f"{self.inspect_list}")
                        self.mail.mail_text_sender(self.DIR_WATCH,self.inspect_list)
            self.previous_files = current_files


class DBConnect:
    def __init__(self):
        self.client = MongoClient('mongodb://localhost:27017')
        self.db = self.client['user']
        self.collection = self.db['userinfo']

    def inData(self,filename, date):
        #date = date_obj = datetime.strptime(date, '%Y-%m-%d').date()
        data = {
            'file' : filename,
            'timestamp' : date
        }
        self.collection.insert_one(data)


if __name__ == "__main__":
    #received_email = ""
    #path = 'uploads'
    #mail = Mail(received_email)
    #mail.mail_attach_sender(path, "trans.xlsx")








    received_email = "hanmin9981@naver.com"
    path = 'uploads'
    MT = Monitering(path,received_email)
    MT.inspect_annotation("#")


